"""
file_extractor.py
------------------
Pulls raw text out of an uploaded resume file. Supports PDF, DOCX, TXT.
All extraction happens in-memory / on the temp path passed in — callers
are responsible for cleaning up the temp file afterwards
(see utils/file_cleanup.py).
"""

import pdfplumber
from docx import Document


class FileExtractionError(Exception):
    """Raised when text cannot be extracted from the uploaded file."""


def extract_text(file_path: str, extension: str) -> str:
    extension = extension.lower()

    if extension == "pdf":
        return _extract_from_pdf(file_path)
    if extension == "docx":
        return _extract_from_docx(file_path)
    if extension == "txt":
        return _extract_from_txt(file_path)

    raise FileExtractionError(f"Unsupported file type: {extension}")


def _extract_from_pdf(file_path: str) -> str:
    text_chunks = []
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                text_chunks.append(page_text)
    except Exception as exc:
        raise FileExtractionError(f"Could not read PDF: {exc}")

    text = "\n".join(text_chunks).strip()
    if not text:
        raise FileExtractionError(
            "No readable text found in PDF. It may be a scanned/image-only file."
        )
    return text


def _extract_from_docx(file_path: str) -> str:
    try:
        doc = Document(file_path)
    except Exception as exc:
        raise FileExtractionError(f"Could not read DOCX: {exc}")

    parts = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also pull text out of tables, since many resumes use table layouts
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())

    text = "\n".join(parts).strip()
    if not text:
        raise FileExtractionError("No readable text found in DOCX file.")
    return text


def _extract_from_txt(file_path: str) -> str:
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read().strip()
    except Exception as exc:
        raise FileExtractionError(f"Could not read TXT file: {exc}")

    if not text:
        raise FileExtractionError("TXT file is empty.")
    return text
